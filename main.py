import streamlit as st
import anthropic
import base64
import webvtt
import json
import re
import os
import csv
import io
import zipfile
from datetime import datetime
from docx import Document as DocxDocument

# --- SLIDE PIPELINE HELPERS (additive — does not affect MOGRT path) ---
NUMBER_WORDS = {
    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10,
    'eleven': 11, 'twelve': 12, 'thirteen': 13, 'fourteen': 14, 'fifteen': 15,
    'sixteen': 16, 'seventeen': 17, 'eighteen': 18, 'nineteen': 19, 'twenty': 20,
}
SLIDE_RE = re.compile(
    r'slide\s*[#:\-\.]?\s*(\d{1,3}|' + '|'.join(NUMBER_WORDS.keys()) + r')\b',
    re.IGNORECASE,
)
TIMECODE_RE = re.compile(r'\b(\d{1,2}):(\d{2})(?::(\d{2}))?(?:[.,](\d+))?\b')

def _tc_match_to_seconds(m):
    g1, g2, g3, g4 = m.group(1), m.group(2), m.group(3), m.group(4)
    if g3 is not None:
        secs = int(g1) * 3600 + int(g2) * 60 + int(g3)
    else:
        secs = int(g1) * 60 + int(g2)
    if g4:
        secs += float("0." + g4)
    return float(secs)

def _slide_token_to_int(tok):
    tok = tok.lower()
    return int(tok) if tok.isdigit() else NUMBER_WORDS[tok]

def _walk_paragraphs(docx_obj):
    for p in docx_obj.paragraphs:
        yield p.text
    for table in docx_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text

def parse_slide_markers_from_docx(file_bytes):
    """Walk a Descript .docx export. Pair each 'slide N' token with the most
    recently seen timecode. Returns sorted, deduped [(seconds, slide_num)]."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    last_tc = None
    markers = []
    for text in _walk_paragraphs(doc):
        if not text or not text.strip():
            continue
        for m in TIMECODE_RE.finditer(text):
            last_tc = _tc_match_to_seconds(m)
        for m in SLIDE_RE.finditer(text):
            num = _slide_token_to_int(m.group(1))
            if last_tc is not None:
                markers.append((last_tc, num))
    seen = set()
    out = []
    for t, n in sorted(markers):
        key = (round(t, 2), n)
        if key in seen:
            continue
        seen.add(key)
        out.append((t, n))
    return out

def extract_slide_images(zip_bytes):
    """Read slide image bytes from an uploaded zip.
    Returns {slide_num: (filename, image_bytes)}. Filenames are kept as-is for
    re-bundling into the output zip's slides/ subfolder."""
    images = {}
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for name in zf.namelist():
            if name.endswith('/') or '__MACOSX' in name:
                continue
            base = os.path.basename(name)
            if not base or base.startswith('.'):
                continue
            ext = os.path.splitext(base)[1].lower()
            if ext not in ('.jpg', '.jpeg', '.png'):
                continue
            stem = os.path.splitext(base)[0]
            num_match = re.search(r'(\d+)', stem)
            if num_match:
                with zf.open(name) as src:
                    images[int(num_match.group(1))] = (base, src.read())
    return images

MAX_SLIDE_DURATION = 12.0  # Hard cap per slide (seconds).

# Cheapest/fastest tier - chosen over Opus/Sonnet since this tool runs a few
# times a day per person; swap to "claude-opus-5" or "claude-sonnet-5" for
# higher quality if Haiku's output isn't good enough on real transcripts.
CLAUDE_MODEL = "claude-haiku-4-5"

# Structured-output schema: guarantees Claude's response is valid JSON matching
# this shape, so no markdown-fence stripping or truncation-continuation dance is
# needed the way it was with Gemini's freeform JSON mode.
GRAPHICS_OUTPUT_SCHEMA = {
    "type": "object",
    "properties": {
        "graphics": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "time_in": {"type": "string"},
                    "mogrt_name": {"type": "string"},
                    "Title_Text": {"type": "string"},
                    "Name": {"type": "string"},
                    "Dropline": {"type": "string"},
                    "Main_Text": {"type": "string"},
                    "bullet-01": {"type": "string"},
                    "bullet-02": {"type": "string"},
                    "bullet-03": {"type": "string"},
                    "bullet-04": {"type": "string"},
                    "bullet-05": {"type": "string"},
                },
                "required": ["time_in", "mogrt_name"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["graphics"],
    "additionalProperties": False,
}

def build_slide_clips(markers, images, default_last_duration=MAX_SLIDE_DURATION):
    """Combine markers + images into slide clip dicts ready for JSX serialization.
    Each slide is capped at MAX_SLIDE_DURATION so the underlying video can breathe
    between slides. Final overlap-avoidance vs MOGRTs happens at JSX runtime."""
    if not markers:
        return [], []
    clips = []
    warnings = []
    for i, (t, n) in enumerate(markers):
        if n not in images:
            warnings.append(f"Marker for slide {n} at {t:.2f}s has no matching image in zip.")
            continue
        if i < len(markers) - 1:
            gap = markers[i + 1][0] - t
            dur = max(0.5, min(MAX_SLIDE_DURATION, gap))
        else:
            dur = MAX_SLIDE_DURATION
        filename, _bytes = images[n]
        clips.append({
            "slide_num": n,
            "time_in": t,
            "duration": dur,
            "filename": filename,  # bare name, e.g. "Slide1.jpg" — JSX resolves relative to its own folder
        })
    used_nums = {n for _, n in markers}
    for img_num in sorted(images.keys()):
        if img_num not in used_nums:
            warnings.append(f"Image for slide {img_num} present but no marker references it.")
    return clips, warnings

def generate_slides_csv(slide_clips):
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Slide #", "Time In (s)", "Duration (s)", "Bundled Filename"])
    for s in slide_clips:
        writer.writerow([s["slide_num"], f"{s['time_in']:.3f}", f"{s['duration']:.3f}", s["filename"]])
    return output.getvalue()

# --- PAGE CONFIG ---
st.set_page_config(page_title="WebMD Spotlight Buddy V1.9.2", layout="wide")

# --- SIDEBAR (Logo & Settings) ---
# Key comes from Streamlit secrets (set in .streamlit/secrets.toml locally, or
# the app's Settings -> Secrets on Streamlit Cloud) - never hardcoded in source,
# since this repo is public. Producers never see a key field when it's configured;
# the manual field is just a fallback for local dev/troubleshooting.
try:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
except st.errors.StreamlitSecretNotFoundError:
    api_key = ""

with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/4/42/WebMD_logo.png", width=150)
    st.header("Settings")
    if not api_key:
        api_key = st.text_input("Anthropic API Key", type="password")

# --- MAIN TITLE ---
st.title("WebMD Spotlight Buddy V1.9.2")
st.markdown("Automated Adobe Premiere Pro Script Generator (Direct JSX Injection)")

# --- HELPER FUNCTION: Convert Time to Seconds ---
def time_str_to_seconds(time_val):
    if isinstance(time_val, (int, float)):
        return float(time_val)
    
    time_str = str(time_val).strip().replace(",", ".")
    try:
        match = re.match(r"(?:(\d{1,2}):)?(\d{1,2}):(\d{2})(?:\.(\d*))?", time_str)
        if match:
            h = match.group(1) or 0
            m = match.group(2) or 0
            s = match.group(3) or 0
            ms = match.group(4) or 0
            ms_val = float(f"0.{ms}") if ms else 0.0
            return int(h) * 3600 + int(m) * 60 + int(s) + ms_val
        return float(time_str)
    except:
        return 0.0

# --- HELPER FUNCTION: Generate CSV String ---
def generate_manager_csv(cleaned_data):
    output = io.StringIO()
    writer = csv.writer(output)
    
    headers = [
        "Timecode", "Graphic Type", "Title Text", "Name", 
        "Dropline", "Main Quote / Banner Text", 
        "Bullet 1", "Bullet 2", "Bullet 3", "Bullet 4", "Bullet 5"
    ]
    writer.writerow(headers)
    
    for g in cleaned_data:
        p = g.get('parameters', {})
        row = [
            g.get('raw_time', ''),
            g.get('mogrt_name', ''),
            p.get('Title_Text', ''),
            p.get('Name', ''),
            p.get('Dropline', ''),
            p.get('Main_Text', ''),
            p.get('bullet-01', ''),
            p.get('bullet-02', ''),
            p.get('bullet-03', ''),
            p.get('bullet-04', ''),
            p.get('bullet-05', '')
        ]
        clean_row = [str(item).replace('\n', ' | ') if item else '-' for item in row]
        writer.writerow(clean_row)
        
    return output.getvalue()

# --- MAIN LOGIC ---
col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("1. Upload Transcript (.vtt)", type=["vtt"])
with col2:
    uploaded_pdfs = st.file_uploader("2. Optional: Upload Context PDF(s) (Objectives/Bios/Slides)", type=["pdf"], accept_multiple_files=True)

st.markdown("---")
st.markdown("### Optional: Automated Slide Syncing")
st.caption("Upload a Descript .docx export with `slide N` markers and a .zip of slide JPGs (e.g. `Slide1.jpg`). Leave blank to skip.")
col3, col4 = st.columns(2)
with col3:
    uploaded_slide_doc = st.file_uploader("3. Slide Markers (Descript .docx)", type=["docx"])
with col4:
    uploaded_slide_zip = st.file_uploader("4. Slide Images (.zip of JPGs)", type=["zip"])

if uploaded_file and api_key:
    client = anthropic.Anthropic(api_key=api_key, timeout=600.0, max_retries=3)

    # 1. PARSE VTT
    with open("temp.vtt", "wb") as f: f.write(uploaded_file.getbuffer())
    vtt = webvtt.read("temp.vtt")
    transcript_text = "".join([f"[{c.start}-->{c.end}] {c.text}\n" for c in vtt])
    
    final_timestamp = vtt[-1].end if vtt else "the end of the video"

    # 2. CLAUDE PROMPT
    system_prompt = f"""
    You are an expert Medical Video Editor Assistant. Select graphics based on WebMD rules.
    You will be provided with a VIDEO TRANSCRIPT and an optional CONTEXT PDF.
    
    TEXT GUIDELINES:
    You do NOT need to be strictly verbatim. Create useful, engaging text graphics that guide the learner. 
    Clean up typos, fix spoken grammar, and summarize clearly, but ALWAYS maintain clinical accuracy.
    Do NOT use quotation marks when summarizing information, only use them when quoting verbatim from transcript.
    
    GRAPHIC TYPES & EXACT FIELD RULES (CRITICAL: NEVER alter the 'mogrt_name' ID):

    1. Speaker Intro: 'EDU-GFX-03-SPLIT NAME-HD' (Fields: Name, Dropline).
       - RULE: Create one for EVERY unique speaker/guest. Place at their very first sentence.
       - FORMAT Name: Max 15 chars per line, max 3 lines. Insert '\\n' manually.
         Example: "Melinda J.\\nGooderham,\\nMD, MSc, FRCPC"
       - FORMAT Dropline: Pull credentials from PDF. Max 40 chars per line, max 4 lines. Insert '\\n' manually.
         Example: "Assistant Professor, Queens University\\nMedical Director, SKiN Centre for\\nDermatology\\nPeterborough, Ontario, Canada"

    2. Short Quote (Small): 'EDU-GFX-04-SPLIT-QUOTE-HD' (Field: Main_Text).
       - RULE: Use for short, punchy quotes (< 50 chars). Insert '\\n' every ~25 characters.

    3. Full Screen Quote (Long): 'EDU-GFX-07-FS-HD' (Field: Main_Text).
       - RULE: Use for long quotes (> 90 chars).

    4. Lists: 'EDU-GFX-07-FS Bullet Point-HD' (Fields: Title_Text, bullet-01, bullet-02, bullet-03, bullet-04, bullet-05).
       - RULE: You MUST provide a 'Title_Text' that summarizes the list (e.g., "Key Symptoms").
       - RULE: Always try to provide 5 bullets. Summarize or split concepts to fill them out.

    BANNED GRAPHICS: Do NOT generate 'EDU-GFX-02-TITLE-HD' (Program Title) or 'EDU-GFX-05-BANNER HD' (Banner Quote) under any circumstances. These graphic types are retired and must never appear in your output, even if the content seems to fit them.

    CRITICAL INSTRUCTIONS:
    - ANTI-LAZINESS: The provided transcript ends at exactly {final_timestamp}. You MUST process the ENTIRE transcript from start to finish. Ensure there is a graphic every ~60 seconds all the way up to {final_timestamp}. Do NOT stop early.
    - PACING CAP (HARD RULE): Target roughly ONE graphic per 60 seconds — not more. Before adding any graphic, check the time_in of the PREVIOUS graphic you placed (of ANY type, including a Speaker Intro): if it is less than 50 seconds earlier, do NOT add this one — skip it, even if the moment is quote-worthy. The only exception is the Speaker Intro graphic itself, which may always be placed at a speaker's first sentence regardless of spacing — but the NEXT graphic after it still must wait at least 50 seconds from the Speaker Intro's own time_in. After drafting the full list, re-check every consecutive pair of timestamps and delete entries that violate the 50-second minimum, keeping only the stronger of the two.
    - TIMESTAMPS: The 'time_in' field MUST be formatted as a string (e.g., "00:19:23"). Do NOT use decimals.
    
    OUTPUT FORMAT (a "graphics" array):
    {{
        "graphics": [
            {{
                "time_in": "00:01:30",
                "mogrt_name": "EDU-GFX-03-SPLIT NAME-HD",
                "Name": "Melinda J.\\nGooderham,\\nMD, MSc, FRCPC",
                "Dropline": "Assistant Professor, Queens University\\nMedical Director, SKiN Centre for\\nDermatology\\nPeterborough, Ontario, Canada"
            }},
            {{
                "time_in": "00:03:15",
                "mogrt_name": "EDU-GFX-04-SPLIT-QUOTE-HD",
                "Main_Text": "This changes\\nhow we treat\\nchronic cases"
            }}
        ]
    }}
    """

    if st.button("Generate Premiere Script", type="primary"):
        with st.spinner(f"Claude is analyzing the program up to {final_timestamp}. This may take a minute for long videos..."):
            try:
                # --- GENERATION ---
                content_blocks = []
                for pdf_file in uploaded_pdfs:
                    pdf_b64 = base64.standard_b64encode(pdf_file.getvalue()).decode("utf-8")
                    content_blocks.append({
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64},
                    })

                transcript_note = (
                    "\n\nRefer to the PDF(s) above for Learning Objectives, speaker credentials, "
                    "slide content, and accurate clinical data."
                    if uploaded_pdfs else ""
                )
                content_blocks.append({"type": "text", "text": f"VIDEO TRANSCRIPT:\n{transcript_text}{transcript_note}"})

                # Structured output guarantees valid JSON matching GRAPHICS_OUTPUT_SCHEMA, so there's
                # no markdown-fence stripping or truncation-continuation loop needed here.
                with client.messages.stream(
                    model=CLAUDE_MODEL,
                    max_tokens=64000,
                    system=system_prompt,
                    output_config={"format": {"type": "json_schema", "schema": GRAPHICS_OUTPUT_SCHEMA}},
                    messages=[{"role": "user", "content": content_blocks}],
                ) as stream:
                    response = stream.get_final_message()

                if response.stop_reason == "max_tokens":
                    raise RuntimeError(
                        "Claude's response hit the output length limit before finishing. "
                        "Try a shorter transcript or fewer/shorter context PDFs."
                    )

                full_response_text = next(b.text for b in response.content if b.type == "text")
                raw_data = json.loads(full_response_text)["graphics"]

                with st.expander("🔍 View Raw JSON from Claude", expanded=False):
                    st.json(raw_data)
                
                cleaned_data = []
                seen_texts = set()
                
                for item in raw_data:
                    m_name = item.get("mogrt_name") or item.get("graphic") or "UNKNOWN"
                    t_raw = item.get("time_in") or item.get("time") or "00:00:00"
                    t_in_seconds = time_str_to_seconds(t_raw)
                    
                    params = {}
                    for k, v in item.items():
                        if k not in ["time_in", "time", "mogrt_name", "graphic"]:
                            params[k] = v
                    
                    txt_check = params.get("Main_Text", "")
                    if txt_check and txt_check in seen_texts: continue
                    if txt_check: seen_texts.add(txt_check)
                    
                    cleaned_data.append({
                        "raw_time": t_raw,
                        "time_in": t_in_seconds,
                        "mogrt_name": m_name,
                        "parameters": params
                    })

                timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")

                # --- SLIDE PIPELINE (parallel to MOGRT pipeline; does not touch cleaned_data) ---
                slide_clips = []
                slide_warnings = []
                slide_images = {}  # {slide_num: (filename, bytes)} — bundled into output zip below
                if uploaded_slide_doc and uploaded_slide_zip:
                    try:
                        markers = parse_slide_markers_from_docx(uploaded_slide_doc.getvalue())
                        slide_images = extract_slide_images(uploaded_slide_zip.getvalue())
                        slide_clips, slide_warnings = build_slide_clips(markers, slide_images)

                        with st.expander(f"🖼️ Slide Pipeline: {len(slide_clips)} slides paired", expanded=False):
                            st.write(f"**Markers parsed:** {len(markers)}")
                            st.write(f"**Images extracted:** {len(slide_images)}")
                            st.caption("Slide JPGs will be bundled into the output zip under `slides/`. The .jsx resolves them relative to wherever the producer unzips the package.")
                            if slide_warnings:
                                for w in slide_warnings:
                                    st.warning(w)
                            st.json(slide_clips)
                    except Exception as slide_err:
                        st.warning(f"Slide pipeline skipped due to error (MOGRT pipeline unaffected): {slide_err}")
                        slide_clips = []
                        slide_images = {}
                elif uploaded_slide_doc or uploaded_slide_zip:
                    st.info("Slide pipeline needs BOTH the .docx markers and the .zip of images. Skipping slides.")

                csv_content = generate_manager_csv(cleaned_data)
                slides_csv_content = generate_slides_csv(slide_clips) if slide_clips else ""
                
                jsx_content = """
                var seqName = "WebMD_Spotlight_%s";
                var project = app.project;
                
                var logFile = File(Folder.desktop + "/mogrt_debug.txt");
                logFile.open("w");
                function log(msg) { logFile.writeln(msg); }
                log("RUN START: " + new Date().toString());

                var sequence = project.createNewSequence(seqName, "ID_1");
                var track = sequence.videoTracks[1];

                var graphicsList = %s;
                var slideList = %s;

                for (var i = 0; i < graphicsList.length; i++) {
                    var g = graphicsList[i];
                    log("\\n--- Graphic " + (i+1) + ": " + g.mogrt_name + " ---");
                    
                    if (g.mogrt_name == "UNKNOWN") {
                        log("SKIPPING: Claude failed to name this graphic.");
                        continue;
                    }

                    var mogrtItem = findItemStrict(project.rootItem, g.mogrt_name);
                    
                    if (mogrtItem) {
                        log("Found MOGRT: " + mogrtItem.name);
                        var timeObj = new Time();
                        timeObj.seconds = g.time_in;
                        track.overwriteClip(mogrtItem, timeObj.ticks);
                        
                        var newClip = findClipAtTime(track, g.time_in);
                        if (newClip) {
                            
                            var currentDuration = newClip.end.seconds - newClip.start.seconds;
                            var expandedTime = new Time();
                            expandedTime.seconds = newClip.start.seconds + (currentDuration * 2);
                            newClip.end = expandedTime;

                            var mgt = newClip.getMGTComponent();
                            if (mgt) {
                                var props = mgt.properties;
                                
                                var allProps = [];
                                for(var p=0; p<props.numItems; p++) { allProps.push(props[p].displayName); }
                                log("SPY REPORT: [" + allProps.join(", ") + "]");

                                for (var key in g.parameters) {
                                    var val = g.parameters[key];
                                    var success = false;
                                    
                                    for (var j=0; j<props.numItems; j++) {
                                        if (props[j].displayName == key) {
                                            props[j].setValue(val); 
                                            log("EXACT MATCH: '" + key + "' | INJECTED: '" + val.replace(/\\n/g, ' [LINE BREAK] ') + "'");
                                            success=true; break;
                                        }
                                    }
                                    
                                    if (!success) {
                                        var targets = [];
                                        
                                        if (key == "Main_Text") targets = ["TEXT", "Main Text", "Source Text", "Quote", "Text", "Banner Quote", "Full Screen Quote"];
                                        if (key == "Title_Text") targets = ["TITLE", "Title", "Header", "Headline", "Text", "TEXT"];
                                        
                                        if (key.indexOf("bullet") > -1 || key.indexOf("Bullet") > -1) {
                                            var num = key.replace(/[^0-9]/g, ''); 
                                            var padded = (num.length < 2) ? "0"+num : num; 
                                            targets = ["bullet-" + padded, "bullet-" + num, "Bullet Point " + num, "Point " + num];
                                        }
                                        
                                        if (key == "Name") targets = ["NAME", "Speaker Name", "Guest Name", "Split Screen Name", "TEXT"];
                                        if (key == "Dropline") targets = ["DROPLINE", "Speaker Title", "Role", "Job Title"];

                                        for (var t=0; t<targets.length; t++) {
                                            for (var j=0; j<props.numItems; j++) {
                                                if (props[j].displayName.toUpperCase().indexOf(targets[t].toUpperCase()) > -1) {
                                                    props[j].setValue(val);
                                                    log("MATCHED: '" + props[j].displayName + "' using target '" + targets[t] + "' | INJECTED: '" + val.replace(/\\n/g, ' [LINE BREAK] ') + "'");
                                                    success = true; break;
                                                }
                                            }
                                            if (success) break;
                                        }
                                    }
                                    if (!success) log("FAILURE: Could not match key '" + key + "' to any property.");
                                }
                            }
                        }
                    } else {
                        log("CRITICAL ERROR: MOGRT NOT FOUND. Check Bin for: " + g.mogrt_name);
                    }
                }
                
                // === SLIDE PIPELINE (parallel to MOGRT logic; uses videoTracks[0]) ===
                log("\\n=== SLIDE PIPELINE: " + slideList.length + " slide(s) ===");
                if (slideList.length > 0) {
                    var slideTrack = sequence.videoTracks[0];

                    // Resolve slide folder relative to THIS script's location.
                    // Producer unzips the package anywhere; slides/ sits next to the .jsx.
                    var scriptFile = new File($.fileName);
                    var slideFolder = scriptFile.parent.fsName + "/slides/";
                    log("Slide folder (relative to script): " + slideFolder);

                    // Build MOGRT busy intervals from videoTracks[1] so slides never overlap a graphic.
                    var mogrtBusy = [];
                    var mogrtTrackRef = sequence.videoTracks[1];
                    for (var bi = 0; bi < mogrtTrackRef.clips.numItems; bi++) {
                        var bc = mogrtTrackRef.clips[bi];
                        mogrtBusy.push({ start: bc.start.seconds, end: bc.end.seconds });
                    }
                    log("MOGRT busy intervals: " + mogrtBusy.length);

                    // Snapshot existing root bin children by name so we can identify newly imported items.
                    var preExisting = {};
                    for (var pi = 0; pi < project.rootItem.children.numItems; pi++) {
                        preExisting[project.rootItem.children[pi].name] = true;
                    }

                    // Build absolute paths from the script's folder + bundled filename, then import all at once.
                    var slidePaths = [];
                    for (var si = 0; si < slideList.length; si++) {
                        var f = new File(slideFolder + slideList[si].filename);
                        if (!f.exists) {
                            log("MISSING SLIDE FILE: " + f.fsName);
                            continue;
                        }
                        slidePaths.push(f.fsName);
                    }
                    try {
                        if (slidePaths.length > 0) project.importFiles(slidePaths, false, project.rootItem, false);
                    } catch (impErr) {
                        log("SLIDE IMPORT ERROR: " + impErr.toString());
                    }

                    // Map newly imported items by name.
                    var slideItemMap = {};
                    for (var pi2 = 0; pi2 < project.rootItem.children.numItems; pi2++) {
                        var newItem = project.rootItem.children[pi2];
                        if (!preExisting[newItem.name]) {
                            slideItemMap[newItem.name] = newItem;
                        }
                    }

                    for (var si = 0; si < slideList.length; si++) {
                        var s = slideList[si];
                        var fname = s.filename;
                        var stem = fname.replace(/\\.[^.]+$/, '');
                        var item = slideItemMap[fname] || slideItemMap[stem];
                        if (!item) {
                            for (var k = 0; k < project.rootItem.children.numItems; k++) {
                                var c = project.rootItem.children[k];
                                if (c.name == fname || c.name == stem) { item = c; break; }
                            }
                        }
                        if (!item) { log("SLIDE MISSING IMPORT: " + fname); continue; }

                        // Compute desired window, then trim against MOGRT busy intervals.
                        var startSec = s.time_in;
                        var endSec = s.time_in + s.duration;

                        // Push start past any MOGRT covering the slide's start.
                        for (var bj = 0; bj < mogrtBusy.length; bj++) {
                            if (startSec >= mogrtBusy[bj].start && startSec < mogrtBusy[bj].end) {
                                startSec = mogrtBusy[bj].end;
                            }
                        }
                        // Trim end to the earliest MOGRT.start that falls inside the slide window.
                        for (var bk = 0; bk < mogrtBusy.length; bk++) {
                            if (mogrtBusy[bk].start > startSec && mogrtBusy[bk].start < endSec) {
                                endSec = mogrtBusy[bk].start;
                            }
                        }
                        // Re-cap at 12s in case start moved.
                        if (endSec - startSec > 12) endSec = startSec + 12;

                        if (endSec - startSec < 0.5) {
                            log("SLIDE " + s.slide_num + " skipped (would fully overlap MOGRT): " + fname);
                            continue;
                        }

                        var tIn = new Time(); tIn.seconds = startSec;
                        try {
                            slideTrack.overwriteClip(item, tIn.ticks);
                        } catch (slErr) {
                            log("SLIDE OVERWRITE ERROR slide " + s.slide_num + ": " + slErr.toString());
                            continue;
                        }

                        for (var c2 = 0; c2 < slideTrack.clips.numItems; c2++) {
                            var cl = slideTrack.clips[c2];
                            if (Math.abs(cl.start.seconds - startSec) < 0.5) {
                                var endT = new Time(); endT.seconds = endSec;
                                cl.end = endT;
                                var adj = (Math.abs(startSec - s.time_in) > 0.01) ? " (start shifted from " + s.time_in.toFixed(2) + ")" : "";
                                log("SLIDE " + s.slide_num + " @" + startSec.toFixed(2) + "s, dur=" + (endSec - startSec).toFixed(2) + "s" + adj);
                                break;
                            }
                        }
                    }
                }

                logFile.close();
                alert("Done! Check mogrt_debug.txt on Desktop.");

                function findItemStrict(folder, name) {
                    for (var i = 0; i < folder.children.numItems; i++) {
                        var item = folder.children[i];
                        if (item.type == ProjectItemType.BIN) {
                            var found = findItemStrict(item, name);
                            if (found) return found;
                        } else {
                            if (item.name.indexOf(name) >= 0) {
                                if (name == "EDU-GFX-07-FS-HD" && item.name.indexOf("Bullet") > -1) continue;
                                return item;
                            }
                        }
                    }
                    return null;
                }

                function findClipAtTime(track, timeSeconds) {
                    for (var i=0; i<track.clips.numItems; i++) {
                        if (Math.abs(track.clips[i].start.seconds - timeSeconds) < 0.5) return track.clips[i];
                    }
                    return null;
                }
                """ % (timestamp, json.dumps(cleaned_data), json.dumps(slide_clips))

                # CREATE THE ZIP FILE
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr(f"WebMD_Spotlight_{timestamp}.jsx", jsx_content.encode("utf-8"))
                    zf.writestr(f"WebMD_Spotlight_Overview_{timestamp}.csv", csv_content.encode("utf-8"))
                    if slides_csv_content:
                        zf.writestr(f"WebMD_Slides_Overview_{timestamp}.csv", slides_csv_content.encode("utf-8"))
                    # Bundle slide JPGs alongside the .jsx so the script can find them via a relative path.
                    for _num, (fname, fbytes) in slide_images.items():
                        zf.writestr(f"slides/{fname}", fbytes)
                    if slide_clips:
                        zf.writestr(
                            "README_FOR_EDITOR.txt",
                            (
                                "Spotlight Buddy package\n"
                                "=======================\n\n"
                                "1. Unzip this entire folder anywhere on your Mac (Desktop is fine).\n"
                                "2. Open your Premiere project.\n"
                                "3. Double-click the .jsx file (or run it via File > Scripts > Run Script File).\n\n"
                                "IMPORTANT: Keep the .jsx file and the 'slides' folder TOGETHER in the same\n"
                                "folder. The script reads slide images from ./slides/ relative to itself.\n"
                            ).encode("utf-8"),
                        )
                
                zip_bytes = zip_buffer.getvalue()

                st.success("Script and CSV Generated Successfully!")
                st.download_button("📥 Download Spotlight Package (.zip)", zip_bytes, f"WebMD_Spotlight_Package_{timestamp}.zip", "application/zip", use_container_width=True)

            except json.JSONDecodeError as e:
                st.error(f"Failed to parse JSON (the AI likely hit the length limit or hallucinated structure): {e}")
            except anthropic.AuthenticationError as e:
                st.error(f"Authentication error from Claude: {e.message}. Check that your Anthropic API key is valid.")
            except anthropic.PermissionDeniedError as e:
                st.error(f"Claude rejected the request due to a permissions issue: {e.message}")
            except anthropic.RateLimitError:
                st.error("Claude rate-limited this request (429). Wait a moment and try again.")
            except anthropic.APITimeoutError:
                st.error("Claude took too long to respond, even with the extended timeout. This usually means the transcript + PDFs are large, or the API is under heavy load. Try again, or split the transcript into a shorter section.")
            except anthropic.APIStatusError as e:
                st.error(f"Claude server error ({e.status_code}): {e.message}. Try again in a few minutes.")
            except anthropic.APIConnectionError:
                st.error("Network error connecting to Claude's API. Check your internet connection and try again.")
            except Exception as e:
                st.error(f"Error: {e}")